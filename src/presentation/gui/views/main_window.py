"""
Fenetre principale - Version corrigee sans bugs d'ascenseur
"""
import customtkinter as ctk
import asyncio
from typing import Optional, Tuple, Dict, Any
import threading
import json
import os
from datetime import datetime
from tkinter import filedialog, messagebox
import tkinter as tk

# Import conditionnel pour Word
try:
    from docx import Document
    from docx.shared import Inches
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    print("ATTENTION: python-docx non disponible. Export Word desactive.")

from src.presentation.gui.viewmodels.main_viewmodel import MainViewModel, ViewState, HighlightViewModel
from src.presentation.gui.components.progress_circle import ProgressCircle
from src.presentation.gui.components.highlight_card import HighlightGrid
from src.presentation.gui.components.zone_picker import ZonePickerButton


class IntegratedEditPanel(ctk.CTkFrame):
    """Panneau d'edition integre en bas de l'interface."""
    
    def __init__(self, parent, on_save=None, on_cancel=None, on_delete=None):
        super().__init__(parent)
        self.configure(fg_color="#2a2a2a", corner_radius=10)
        self.on_save = on_save
        self.on_cancel = on_cancel
        self.on_delete = on_delete
        self.current_highlight = None
        
        # CORRECTION BUG MODIF: Variables pour l'etat initial
        self.initial_text = ""
        self.initial_name = ""
        
        self._create_content()
        # Le placement est gere par le parent avec grid()
        
        # Variables pour l'historique d'annulation
        self._text_history = []
        self._name_history = []
        self._max_history = 50
    
    def _on_key_press_name(self, event):
        """Sauvegarde dans l'historique avant modification du nom."""
        # CORRECTION CTRL+Z: Sauvegarder seulement pour les vraies modifications
        if event.keysym not in ['Control_L', 'Control_R', 'Shift_L', 'Shift_R', 'Alt_L', 'Alt_R', 
                               'Left', 'Right', 'Up', 'Down', 'Home', 'End', 'Tab']:
            current_value = self.name_entry.get()
            if not self._name_history or (self._name_history and self._name_history[-1] != current_value):
                if len(self._name_history) >= self._max_history:
                    self._name_history.pop(0)
                self._name_history.append(current_value)
                print(f"DEBUG: Historique nom sauvegarde: '{current_value}' (touche: {event.keysym})")
    
    def _on_key_press_text(self, event):
        """Sauvegarde dans l'historique avant modification du texte."""
        # CORRECTION CTRL+Z: Sauvegarder seulement pour les vraies modifications
        if event.keysym not in ['Control_L', 'Control_R', 'Shift_L', 'Shift_R', 'Alt_L', 'Alt_R',
                               'Left', 'Right', 'Up', 'Down', 'Home', 'End', 'Tab', 'Prior', 'Next']:
            current_value = self.text_editor.get("1.0", "end-1c")
            if not self._text_history or (self._text_history and self._text_history[-1] != current_value):
                if len(self._text_history) >= self._max_history:
                    self._text_history.pop(0)
                self._text_history.append(current_value)
                print(f"DEBUG: Historique texte sauvegarde: '{current_value[:30]}...' (touche: {event.keysym})")
    
    def _on_focus_out(self, event):
        """Sauvegarde automatique quand on perd le focus (clic ailleurs)."""
        self._save_changes()
    
    def _on_enter_pressed(self, event):
        """Sauvegarde automatique quand on appuie sur Entree."""
        self._save_changes()
        return "break"  # Empeche le comportement par defaut
    
    def _on_undo_pressed(self, event):
        """Gere l'annulation (Ctrl+Z) - VERSION CORRIGEE RENFORCEE."""
        widget = event.widget
        print(f"DEBUG: Ctrl+Z presse sur widget: {type(widget)}")
        
        # CORRECTION RENFORCEE: Multiples tentatives d'acces aux widgets natifs
        try:
            # Methode 1: Acces direct aux widgets sous-jacents CustomTkinter
            if hasattr(widget, '_entry'):
                tk_widget = widget._entry
                print(f"DEBUG: Widget Entry trouve: {type(tk_widget)}")
                if hasattr(tk_widget, 'edit_undo'):
                    try:
                        tk_widget.edit_undo()
                        print("SUCCESS: Undo natif CTkEntry execute")
                        return "break"
                    except tk.TclError as e:
                        print(f"DEBUG: Undo natif Entry echoue: {e}")
            
            elif hasattr(widget, '_textbox'):
                tk_widget = widget._textbox
                print(f"DEBUG: Widget Textbox trouve: {type(tk_widget)}")
                if hasattr(tk_widget, 'edit_undo'):
                    try:
                        tk_widget.edit_undo()
                        print("SUCCESS: Undo natif CTkTextbox execute")
                        return "break"
                    except tk.TclError as e:
                        print(f"DEBUG: Undo natif Textbox echoue: {e}")
            
            # Methode 2: Si c'est directement un widget Tkinter
            elif hasattr(widget, 'edit_undo'):
                try:
                    widget.edit_undo()
                    print("SUCCESS: Undo natif direct execute")
                    return "break"
                except tk.TclError as e:
                    print(f"DEBUG: Undo natif direct echoue: {e}")
        
        except Exception as e:
            print(f"DEBUG: Erreur acces widgets natifs: {e}")
        
        # METHODE 3: Forcer l'activation de l'undo sur les widgets natifs
        try:
            if hasattr(widget, '_entry'):
                tk_widget = widget._entry
                # Forcer l'activation de l'undo si pas deja fait
                try:
                    tk_widget.configure(undo=True, maxundo=20)
                    tk_widget.edit_undo()
                    print("SUCCESS: Undo force CTkEntry execute")
                    return "break"
                except:
                    pass
            
            elif hasattr(widget, '_textbox'):
                tk_widget = widget._textbox
                # Forcer l'activation de l'undo si pas deja fait
                try:
                    tk_widget.configure(undo=True, maxundo=20)
                    tk_widget.edit_undo()
                    print("SUCCESS: Undo force CTkTextbox execute")
                    return "break"
                except:
                    pass
        
        except Exception as e:
            print(f"DEBUG: Undo force echoue: {e}")
        
        # FALLBACK: Systeme d'undo manuel ameliore
        print("DEBUG: Utilisation du systeme d'undo manuel")
        try:
            # Identifier le widget parent CustomTkinter
            if hasattr(widget, 'master') and hasattr(widget.master, 'master'):
                parent_widget = widget.master.master
                if parent_widget == self.name_entry:
                    print("DEBUG: Undo manuel pour champ nom")
                    if self._name_history and len(self._name_history) >= 1:
                        if len(self._name_history) >= 2:
                            previous_value = self._name_history[-2]
                            self._name_history.pop()
                        else:
                            previous_value = ""
                        
                        self.name_entry.delete(0, 'end')
                        if previous_value:
                            self.name_entry.insert(0, previous_value)
                        print(f"SUCCESS: Undo manuel nom: -> '{previous_value}'")
                        return "break"
                
                elif parent_widget == self.text_editor:
                    print("DEBUG: Undo manuel pour champ texte")
                    if self._text_history and len(self._text_history) >= 1:
                        if len(self._text_history) >= 2:
                            previous_value = self._text_history[-2]
                            self._text_history.pop()
                        else:
                            previous_value = ""
                        
                        self.text_editor.delete("1.0", "end")
                        if previous_value:
                            self.text_editor.insert("1.0", previous_value)
                        print(f"SUCCESS: Undo manuel texte: -> '{previous_value[:30]}...'")
                        return "break"
            
            # Si l'identification par parent ne marche pas, essayer par comparaison directe
            if widget == self.name_entry or str(widget).find('entry') != -1:
                print("DEBUG: Undo manuel pour nom (identification directe)")
                if self._name_history:
                    if len(self._name_history) >= 2:
                        previous_value = self._name_history[-2]
                        self._name_history.pop()
                    else:
                        previous_value = ""
                    
                    self.name_entry.delete(0, 'end')
                    if previous_value:
                        self.name_entry.insert(0, previous_value)
                    print(f"SUCCESS: Undo manuel nom direct: -> '{previous_value}'")
                    return "break"
            
            elif widget == self.text_editor or str(widget).find('text') != -1:
                print("DEBUG: Undo manuel pour texte (identification directe)")
                if self._text_history:
                    if len(self._text_history) >= 2:
                        previous_value = self._text_history[-2]
                        self._text_history.pop()
                    else:
                        previous_value = ""
                    
                    self.text_editor.delete("1.0", "end")
                    if previous_value:
                        self.text_editor.insert("1.0", previous_value)
                    print(f"SUCCESS: Undo manuel texte direct: -> '{previous_value[:30]}...'")
                    return "break"
        
        except Exception as e:
            print(f"ERREUR: Undo manuel echoue: {e}")
        
        print("WARNING: Aucune methode d'undo n'a fonctionne")
        return "break"
    
    def _save_text_to_history(self):
        """Sauvegarde l'etat actuel dans l'historique."""
        if self.current_highlight:
            try:
                # Sauvegarder l'etat du nom
                current_name = self.name_entry.get()
                if not self._name_history or (self._name_history and self._name_history[-1] != current_name):
                    if len(self._name_history) >= self._max_history:
                        self._name_history.pop(0)
                    self._name_history.append(current_name)
                    print(f"DEBUG: Nom sauvegarde dans historique: '{current_name}'")
                
                # Sauvegarder l'etat du texte
                current_text = self.text_editor.get("1.0", "end-1c")
                if not self._text_history or (self._text_history and self._text_history[-1] != current_text):
                    if len(self._text_history) >= self._max_history:
                        self._text_history.pop(0)
                    self._text_history.append(current_text)
                    print(f"DEBUG: Texte sauvegarde dans historique: '{current_text[:30]}...'")
            
            except Exception as e:
                print(f"ERREUR: Sauvegarde historique echouee: {e}")
    
    def _create_content(self):
        """Cree le contenu du panneau d'edition."""
        # Header
        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.pack(fill="x", padx=15, pady=(15, 10))
        
        title_label = ctk.CTkLabel(
            header_frame,
            text="EDITION DU HIGHLIGHT",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color="#ffffff"
        )
        title_label.pack(side="left")
        
        # Corps principal avec 2 colonnes
        main_frame = ctk.CTkFrame(self, fg_color="transparent")
        main_frame.pack(fill="both", expand=True, padx=15, pady=(0, 15))
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_columnconfigure(1, weight=1)
        
        # Colonne gauche - Informations
        left_column = ctk.CTkFrame(main_frame, fg_color="#3a3a3a")
        left_column.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        
        # Page et confiance
        info_frame = ctk.CTkFrame(left_column, fg_color="transparent")
        info_frame.pack(fill="x", padx=15, pady=15)
        
        self.page_label = ctk.CTkLabel(
            info_frame,
            text="Page: -",
            font=ctk.CTkFont(size=12, weight="bold")
        )
        self.page_label.pack(side="left")
        
        self.confidence_label = ctk.CTkLabel(
            info_frame,
            text="Confiance: -",
            font=ctk.CTkFont(size=12)
        )
        self.confidence_label.pack(side="right")
        
        # Champ nom personnalise
        name_label = ctk.CTkLabel(
            left_column,
            text="Nom personnalise:",
            font=ctk.CTkFont(size=12, weight="bold"),
            anchor="w"
        )
        name_label.pack(fill="x", padx=15, pady=(0, 5))
        
        self.name_entry = ctk.CTkEntry(
            left_column,
            placeholder_text="Nom pour ce highlight...",
            font=ctk.CTkFont(size=11)
        )
        self.name_entry.pack(fill="x", padx=15, pady=(0, 15))
        
        # Bindings pour sauvegarde automatique et undo
        self.name_entry.bind("<Return>", self._on_enter_pressed)
        self.name_entry.bind("<Control-z>", self._on_undo_pressed)
        self.name_entry.bind("<FocusOut>", self._on_focus_out)
        self.name_entry.bind("<KeyPress>", self._on_key_press_name)
        
        # CORRECTION CTRL+Z: Forcer l'activation de l'undo sur le widget natif
        try:
            if hasattr(self.name_entry, '_entry'):
                self.name_entry._entry.configure(undo=True, maxundo=20)
                print("DEBUG: Undo active sur name_entry")
        except Exception as e:
            print(f"DEBUG: Erreur activation undo name_entry: {e}")
        
        # Boutons d'action
        buttons_frame = ctk.CTkFrame(left_column, fg_color="transparent")
        buttons_frame.pack(fill="x", padx=15, pady=(0, 15))
        
        clear_btn = ctk.CTkButton(
            buttons_frame,
            text="SUPPRIMER FICHE",
            command=self._delete_highlight,
            fg_color="#cc4444",
            hover_color="#dd5555",
            height=35
        )
        clear_btn.pack(fill="x")
        
        # Colonne droite - Texte
        right_column = ctk.CTkFrame(main_frame, fg_color="#3a3a3a")
        right_column.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        
        text_label = ctk.CTkLabel(
            right_column,
            text="Texte extrait:",
            font=ctk.CTkFont(size=12, weight="bold"),
            anchor="w"
        )
        text_label.pack(fill="x", padx=15, pady=(15, 5))
        
        self.text_editor = ctk.CTkTextbox(
            right_column,
            font=ctk.CTkFont(size=11),
            wrap="word"
        )
        self.text_editor.pack(fill="both", expand=True, padx=15, pady=(0, 15))
        
        # Bindings pour sauvegarde automatique et undo
        self.text_editor.bind("<Return>", self._on_enter_pressed)
        self.text_editor.bind("<Control-z>", self._on_undo_pressed)
        self.text_editor.bind("<FocusOut>", self._on_focus_out)
        self.text_editor.bind("<KeyPress>", self._on_key_press_text)
        
        # CORRECTION CTRL+Z: Forcer l'activation de l'undo sur le widget natif
        try:
            if hasattr(self.text_editor, '_textbox'):
                self.text_editor._textbox.configure(undo=True, maxundo=20)
                print("DEBUG: Undo active sur text_editor")
        except Exception as e:
            print(f"DEBUG: Erreur activation undo text_editor: {e}")
        
        # Afficher un message par defaut
        self._show_default_message()
    
    def _show_default_message(self):
        """Affiche un message par defaut quand aucun highlight n'est selectionne."""
        self.page_label.configure(text="Page: -")
        self.confidence_label.configure(text="Confiance: -")
        self.name_entry.delete(0, 'end')
        self.text_editor.delete("1.0", "end")
        self.text_editor.insert("1.0", "Selectionnez un highlight ci-dessus pour l'afficher ici...")
        self.current_highlight = None
        
        # CORRECTION BUG MODIF: Reinitialiser l'etat initial
        self.initial_text = ""
        self.initial_name = ""
    
    def show_highlight(self, highlight_data):
        """Affiche un highlight pour edition."""
        # Sauvegarder l'etat actuel avant de changer
        self._save_text_to_history()
        
        self.current_highlight = highlight_data.copy()
        
        # Remplir les champs
        page = highlight_data.get('page', '?')
        confidence = highlight_data.get('confidence', 0)
        
        self.page_label.configure(text=f"Page: {page}")
        self.confidence_label.configure(text=f"Confiance: {confidence:.1f}%")
        
        # Nom personnalise
        custom_name = highlight_data.get('custom_name', '')
        self.name_entry.delete(0, 'end')
        if custom_name:
            self.name_entry.insert(0, custom_name)
        
        # Texte
        text = highlight_data.get('text', '')
        self.text_editor.delete("1.0", "end")
        self.text_editor.insert("1.0", text)
        
        # CORRECTION BUG MODIF: Stocker l'etat initial pour comparaison future
        self.initial_name = custom_name
        self.initial_text = text
        print(f"DEBUG: Etat initial sauvegarde - Nom: '{self.initial_name}', Texte: '{self.initial_text[:30]}...'")
        
        # Reinitialiser l'historique pour le nouveau highlight
        self._text_history.clear()
        self._name_history.clear()
    
    def _save_changes(self):
        """Sauvegarde les modifications."""
        if not self.current_highlight:
            return
        
        # Recuperer les nouvelles valeurs
        new_text = self.text_editor.get("1.0", "end-1c")
        new_name = self.name_entry.get().strip()
        
        # CORRECTION BUG MODIF: Verifier s'il y a vraiment eu des modifications
        has_changes = False
        
        # Comparer le texte
        if new_text != self.initial_text:
            has_changes = True
            print(f"DEBUG: Changement texte detecte: '{self.initial_text[:30]}...' -> '{new_text[:30]}...'")
        
        # Comparer le nom
        if new_name != self.initial_name:
            has_changes = True
            print(f"DEBUG: Changement nom detecte: '{self.initial_name}' -> '{new_name}'")
        
        # Mettre a jour les donnees
        self.current_highlight['text'] = new_text
        if new_name:
            self.current_highlight['custom_name'] = new_name
        elif 'custom_name' in self.current_highlight:
            del self.current_highlight['custom_name']
        
        # CORRECTION BUG MODIF: Marquer comme modifie SEULEMENT s'il y a vraiment eu des changements
        if has_changes:
            self.current_highlight['modified'] = True
            self.current_highlight['modified_date'] = datetime.now().isoformat()
            print("DEBUG: Highlight marque comme modifie")
        else:
            print("DEBUG: Aucun changement detecte, pas de modification marquee")
        
        # Callback de sauvegarde
        if self.on_save:
            self.on_save(self.current_highlight)
    
    def _delete_highlight(self):
        """Supprime la fiche highlight actuellement selectionnee."""
        if not self.current_highlight:
            messagebox.showwarning("Aucune selection", "Aucun highlight selectionne a supprimer.")
            return
        
        # Obtenir le nom de la fiche pour la confirmation
        page = self.current_highlight.get('page', '?')
        name = self.current_highlight.get('custom_name', f"Page {page}")
        
        # Demander confirmation
        if messagebox.askyesno(
            "Supprimer le highlight",
            f"Etes-vous sur de vouloir supprimer definitivement :\n\n'{name}'\n\nCette action est irreversible."
        ):
            # Notifier le parent pour supprimer la fiche
            if self.on_delete:
                self.on_delete(self.current_highlight)
            
            # Revenir a l'etat par defaut
            self._show_default_message()


class MainWindow(ctk.CTk):
    """
    Fenetre principale de l'application AllamBik v3 - Version corrigee.
    """
    
    def __init__(self, viewmodel: MainViewModel):
        super().__init__()
        
        self.viewmodel = viewmodel
        self.async_loop = None
        
        # Variables d'etat (AVANT _create_ui())
        self.view_mode = "grid"  # "grid" (2 col) ou "list" (1 col)
        self.current_search = ""
        self.extraction_file_path = None  # Chemin vers le fichier d'extraction
        self.selected_card = None  # Fiche actuellement selectionnee
        
        # Pour gerer les mises a jour
        self._update_queue = []
        
        self._setup_window()
        self._create_ui()
        self._bind_viewmodel()
        self._start_update_loop()
    
    def _setup_window(self):
        """Configure la fenetre principale."""
        self.title("AllamBik v3.0 - Extracteur de Highlights Kindle")
        self.geometry("1400x900")
        self.minsize(1200, 700)
        
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        self.colors = {
            'bg_primary': '#1a1a1a',
            'bg_secondary': '#2a2a2a',
            'bg_tertiary': '#3a3a3a',
            'text_primary': '#ffffff',
            'text_secondary': '#b0b0b0',
            'text_muted': '#808080',
            'accent': '#00aaff',
            'success': '#00ff88',
            'warning': '#ffaa00',
            'error': '#ff4444'
        }
        
        self.configure(fg_color=self.colors['bg_primary'])
    
    def _create_ui(self):
        """Cree l'interface utilisateur."""
        # Container principal
        main_container = ctk.CTkFrame(self, fg_color="transparent")
        main_container.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Configuration de la grille
        main_container.grid_columnconfigure(0, weight=1, minsize=400)
        main_container.grid_columnconfigure(1, weight=2)
        main_container.grid_rowconfigure(0, weight=1)
        
        # Panneau gauche (controles)
        self._create_left_panel(main_container)
        
        # Panneau droit (highlights + edition)
        self._create_right_panel(main_container)
    
    def _create_left_panel(self, parent):
        """Cree le panneau gauche avec controles et progression."""
        left_panel = ctk.CTkFrame(parent, fg_color="transparent")
        left_panel.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        
        # Header
        header_frame = ctk.CTkFrame(
            left_panel,
            fg_color=self.colors['bg_secondary'],
            height=60,
            corner_radius=10
        )
        header_frame.pack(fill="x", pady=(0, 20))
        header_frame.pack_propagate(False)
        
        title_label = ctk.CTkLabel(
            header_frame,
            text="ALLAMBIK v3.0",
            font=ctk.CTkFont(size=20, weight="bold"),
            text_color=self.colors['text_primary']
        )
        title_label.pack(expand=True)
        
        # Section Progression
        progress_section = self._create_section(left_panel, "PROGRESSION")
        self.progress_circle = ProgressCircle(progress_section)
        self.progress_circle.pack(pady=10)
        
        self.progress_label = ctk.CTkLabel(
            progress_section,
            text="Pret pour l'extraction",
            font=ctk.CTkFont(size=12),
            text_color=self.colors['text_secondary']
        )
        self.progress_label.pack(pady=(0, 15))
        
        # Section Controles
        controls_section = self._create_section(left_panel, "CONTROLES")
        
        # Bouton Zone Picker
        self.zone_picker_button = ZonePickerButton(
            controls_section,
            on_zone_selected=self._on_zone_selected,
            text="DEFINIR ZONE DE SCAN",
            fg_color=self.colors['warning'],
            hover_color="#ff8800"
        )
        self.zone_picker_button.pack(fill="x", padx=20, pady=(10, 5))
        
        # Bouton Demarrer
        self.start_button = ctk.CTkButton(
            controls_section,
            text="DEMARRER EXTRACTION",
            font=ctk.CTkFont(size=14, weight="bold"),
            height=45,
            fg_color=self.colors['accent'],
            hover_color="#0088cc",
            command=self._on_start_clicked
        )
        self.start_button.pack(fill="x", padx=20, pady=(5, 5))
        
        # Bouton Arreter
        self.stop_button = ctk.CTkButton(
            controls_section,
            text="ARRETER EXTRACTION",
            font=ctk.CTkFont(size=12, weight="bold"),
            height=40,
            fg_color=self.colors['bg_tertiary'],
            hover_color=self.colors['error'],
            state="disabled",
            command=self._on_stop_clicked
        )
        self.stop_button.pack(fill="x", padx=20, pady=(5, 5))
        
        # Bouton Detection automatique du nombre de pages
        self.detect_button = ctk.CTkButton(
            controls_section,
            text="DETECTER NOMBRE DE PAGES",
            font=ctk.CTkFont(size=12, weight="bold"),
            height=40,
            fg_color="#ff6600",  # Orange
            hover_color="#ff8800",
            command=self._on_detect_pages_clicked
        )
        self.detect_button.pack(fill="x", padx=20, pady=(5, 5))
        
        # Bouton Import JSON pour tests rapides
        self.import_json_button = ctk.CTkButton(
            controls_section,
            text="CHARGER JSON", 
            font=ctk.CTkFont(size=12, weight="bold"),
            height=40,
            fg_color="#8B4513",
            hover_color="#A0522D", 
            command=self._on_import_json_clicked
        )
        self.import_json_button.pack(fill="x", padx=20, pady=(5, 5))
        
        # BOUTON EXPORT WORD - DIRECTEMENT SOUS ARRETER
        self.export_word_button = ctk.CTkButton(
            controls_section,
            text="EXPORTER WORD",
            font=ctk.CTkFont(size=12, weight="bold"),
            height=35,
            fg_color="#2d5a2d",  # Vert fonce pour le distinguer
            hover_color="#3d6a3d",
            command=self._on_export_word_clicked
        )
        self.export_word_button.pack(fill="x", padx=20, pady=(5, 15))
        
        # Section Statistiques
        stats_section = self._create_section(left_panel, "STATISTIQUES")
        
        stats_frame = ctk.CTkFrame(stats_section, fg_color="transparent")
        stats_frame.pack(fill="x", padx=20, pady=(10, 15))
        
        # Creer les stats
        self.stats_labels = {}
        stats = [
            ("pages_scanned", "Pages scannees"),
            ("pages_with_content", "Pages avec contenu"),
            ("highlights_count", "Highlights extraits")
        ]
        
        for i, (key, label) in enumerate(stats):
            stat_container = ctk.CTkFrame(stats_frame, fg_color="transparent")
            stat_container.pack(fill="x", pady=5)
            
            label_widget = ctk.CTkLabel(
                stat_container,
                text=label,
                font=ctk.CTkFont(size=10),
                text_color=self.colors['text_muted'],
                anchor="w"
            )
            label_widget.pack(side="left", fill="x", expand=True)
            
            value_widget = ctk.CTkLabel(
                stat_container,
                text="0",
                font=ctk.CTkFont(size=16, weight="bold"),
                text_color=self.colors['text_primary'],
                anchor="e"
            )
            value_widget.pack(side="right")
            
            self.stats_labels[key] = value_widget
    
    def _create_right_panel(self, parent):
        """Cree le panneau droit avec highlights et panneau d'edition."""
        right_panel = ctk.CTkFrame(parent, fg_color="transparent")
        right_panel.grid(row=0, column=1, sticky="nsew", padx=(10, 0))
        
        # Configuration - highlights en haut, edition en bas
        right_panel.grid_rowconfigure(0, weight=2)  # Highlights prennent plus d'espace
        right_panel.grid_rowconfigure(1, weight=1)  # Panneau d'edition en bas
        right_panel.grid_columnconfigure(0, weight=1)
        
        # Container pour highlights
        highlights_container = ctk.CTkFrame(right_panel, fg_color="transparent")
        highlights_container.grid(row=0, column=0, sticky="nsew", pady=(0, 5))
        
        # Section Highlights
        highlights_section = self._create_section_frame(highlights_container, "HIGHLIGHTS EXTRAITS")
        highlights_section.pack(fill="both", expand=True)
        
        # Header avec compteur et controles SIMPLIFIES
        header_frame = ctk.CTkFrame(highlights_section, fg_color="transparent")
        header_frame.pack(fill="x", padx=15, pady=(15, 5))
        
        # Titre avec compteur
        self.highlights_title = ctk.CTkLabel(
            header_frame,
            text="HIGHLIGHTS EXTRAITS (0)",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color=self.colors['text_primary']
        )
        self.highlights_title.pack(side="left")
        
        # Controles simplifies (SEULEMENT 2 boutons)
        controls_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        controls_frame.pack(side="right")
        
        # Bouton mode liste (1 colonne large)
        self.list_view_button = ctk.CTkButton(
            controls_frame,
            text="LISTE",
            width=60,
            height=35,
            font=ctk.CTkFont(size=12),
            fg_color="#4a4a4a" if self.view_mode != "list" else self.colors['accent'],
            hover_color="#5a5a5a",
            command=self._set_list_view
        )
        self.list_view_button.pack(side="right", padx=(5, 0))
        
        # Bouton mode grille (2 colonnes)
        self.grid_view_button = ctk.CTkButton(
            controls_frame,
            text="GRILLE",
            width=60,
            height=35,
            font=ctk.CTkFont(size=12),
            fg_color="#4a4a4a" if self.view_mode != "grid" else self.colors['accent'],
            hover_color="#5a5a5a",
            command=self._set_grid_view
        )
        self.grid_view_button.pack(side="right", padx=(5, 0))
        
        # Zone de recherche
        search_frame = ctk.CTkFrame(highlights_section, fg_color="transparent")
        search_frame.pack(fill="x", padx=15, pady=(5, 10))
        
        self.search_entry = ctk.CTkEntry(
            search_frame,
            placeholder_text="Rechercher dans les highlights...",
            font=ctk.CTkFont(size=11),
            height=30
        )
        self.search_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))
        self.search_entry.bind("<KeyRelease>", self._on_search_changed)
        
        # Bouton effacer recherche
        clear_search_btn = ctk.CTkButton(
            search_frame,
            text="X",
            width=30,
            height=30,
            fg_color="transparent",
            hover_color="#444444",
            command=self._clear_search
        )
        clear_search_btn.pack(side="right")
        
        # Grille de highlights
        self.highlights_grid = HighlightGrid(
            highlights_section,
            columns=2,  # Demarre en mode grille
            fg_color="transparent",
            on_edit_requested=self._on_edit_requested,
            on_highlight_selected=self._on_highlight_selected  # NOUVEAU: selection simple
        )
        self.highlights_grid.pack(fill="both", expand=True, padx=15, pady=(0, 15))
        
        # Panneau d'edition integre (TOUJOURS VISIBLE)
        self.edit_panel = IntegratedEditPanel(
            right_panel,
            on_save=self._on_edit_saved,
            on_cancel=self._on_edit_cancelled,
            on_delete=self._on_highlight_deleted
        )
        self.edit_panel.grid(row=1, column=0, sticky="ew", pady=(5, 0))
    
    def _create_section(self, parent, title: str) -> ctk.CTkFrame:
        """Cree une section avec titre."""
        section = ctk.CTkFrame(
            parent,
            fg_color=self.colors['bg_secondary'],
            corner_radius=10
        )
        section.pack(fill="x", pady=(0, 20))
        
        title_label = ctk.CTkLabel(
            section,
            text=title,
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color=self.colors['text_primary']
        )
        title_label.pack(pady=(15, 5))
        
        return section
    
    def _create_section_frame(self, parent, title: str) -> ctk.CTkFrame:
        """Cree une section frame pour la grille."""
        section = ctk.CTkFrame(
            parent,
            fg_color=self.colors['bg_secondary'],
            corner_radius=10
        )
        return section
    
    def _bind_viewmodel(self):
        """Lie le ViewModel a la vue."""
        self.viewmodel.on_state_changed = self._on_state_changed
        self.viewmodel.on_progress_changed = self._on_progress_changed
        self.viewmodel.on_highlight_added = self._on_highlight_added
    
    def _start_update_loop(self):
        """Demarre la boucle de mise a jour UI."""
        def update_ui():
            try:
                for update_func in self._update_queue[:]:
                    update_func()
                    self._update_queue.remove(update_func)
            except:
                pass
            self.after(100, update_ui)
        
        self.after(100, update_ui)
    
    def _schedule_update(self, func):
        """Planifie une mise a jour UI thread-safe."""
        self._update_queue.append(func)
    
    # Gestion des vues SIMPLIFIEE
    
    def _set_grid_view(self):
        """Active le mode grille (2 colonnes)."""
        if self.view_mode != "grid":
            self.view_mode = "grid"
            self._update_view_buttons()
            self._recreate_grid(2)
    
    def _set_list_view(self):
        """Active le mode liste (1 colonne large)."""
        if self.view_mode != "list":
            self.view_mode = "list"
            self._update_view_buttons()
            self._recreate_grid(1)
    
    def _update_view_buttons(self):
        """Met a jour l'apparence des boutons de vue."""
        # Bouton grille
        grid_color = self.colors['accent'] if self.view_mode == "grid" else "#4a4a4a"
        self.grid_view_button.configure(fg_color=grid_color)
        
        # Bouton liste
        list_color = self.colors['accent'] if self.view_mode == "list" else "#4a4a4a"
        self.list_view_button.configure(fg_color=list_color)
    
    def _recreate_grid(self, columns):
        """Recree la grille avec le nombre de colonnes specifie - CORRECTION BUG ASCENSEUR."""
        
        # Sauvegarder les donnees
        highlights_data = self.highlights_grid.get_highlights_data()
        
        # Supprimer toutes les cartes existantes de la grille
        for card in self.highlights_grid.cards:
            card.destroy()
        
        # Reinitialiser les listes
        self.highlights_grid.cards.clear()
        self.highlights_grid.highlights_data.clear()
        
        # Mettre a jour le nombre de colonnes de la grille existante
        self.highlights_grid.columns = columns
        
        # Reconfigurer les colonnes de la grille existante
        for i in range(columns):
            self.highlights_grid.grid_columnconfigure(i, weight=1, uniform="column")
        
        # Supprimer la configuration des colonnes supplementaires si on reduit
        if columns < 4:
            for i in range(columns, 4):
                self.highlights_grid.grid_columnconfigure(i, weight=0, uniform="")
        
        # Forcer la mise a jour du scrollable frame
        self.highlights_grid.update_idletasks()
        
        # Attendre un cycle d'evenements puis restaurer les donnees
        self.after_idle(lambda: self._restore_highlights_data(highlights_data))
    
    def _restore_highlights_data(self, highlights_data):
        """Restaure les donnees des highlights dans la grille modifiee."""
        # Sauvegarder la fiche actuellement selectionnee
        selected_highlight = None
        if self.selected_card and self.selected_card.highlight_data:
            selected_highlight = self.selected_card.highlight_data.copy()
        
        # Restaurer toutes les donnees
        for highlight_data in highlights_data:
            self.highlights_grid.add_highlight(highlight_data)
        
        # Restaurer la selection si possible
        if selected_highlight:
            self.after_idle(lambda: self._restore_selection(selected_highlight))
        
        self._update_highlights_count()
        
        # Forcer la mise a jour complete de l'affichage
        self.highlights_grid.update_idletasks()
        self.update_idletasks()
    
    def _restore_selection(self, selected_highlight):
        """Restaure la selection apres changement de mode."""
        for card in self.highlights_grid.cards:
            if (card.highlight_data.get('page') == selected_highlight.get('page') and
                card.highlight_data.get('text', '')[:50] == selected_highlight.get('text', '')[:50]):
                card.set_selected(True)
                self.selected_card = card
                print(f"DEBUG: Selection restauree - Page {selected_highlight.get('page')}")
                break
    
    # Gestion de l'edition integree
    
    def _on_highlight_selected(self, highlight_data):
        """NOUVEAU: Callback quand un highlight est selectionne (simple clic)."""
        # CORRECTION 1: Sauvegarder avant de changer de fiche
        if hasattr(self, 'edit_panel') and self.edit_panel.current_highlight:
            self.edit_panel._save_changes()
        
        # CORRECTION 2: Deselectionner la fiche precedente de maniere robuste
        if self.selected_card:
            try:
                self.selected_card.set_selected(False)
            except:
                pass  # La carte peut avoir ete detruite
        
        # CORRECTION 3: Trouver et selectionner la nouvelle fiche de maniere robuste
        self.selected_card = None
        for card in self.highlights_grid.cards:
            # Criteres multiples pour identifier la bonne carte
            if (card.highlight_data.get('page') == highlight_data.get('page') and
                card.highlight_data.get('text', '')[:50] == highlight_data.get('text', '')[:50]):
                try:
                    card.set_selected(True)
                    self.selected_card = card
                    print(f"DEBUG: Fiche selectionnee - Page {highlight_data.get('page')}")
                    break
                except:
                    print(f"ERREUR: Impossible de selectionner la carte Page {highlight_data.get('page')}")
        
        # Afficher dans le panneau d'edition
        self.edit_panel.show_highlight(highlight_data)
    
    def _on_edit_requested(self, highlight_data):
        """Callback quand l'edition d'un highlight est demandee (double-clic)."""
        self.edit_panel.show_highlight(highlight_data)
    
    def _on_edit_saved(self, updated_data):
        """Callback quand un highlight est sauvegarde apres edition."""
        print(f"DEBUG: Sauvegarde highlight: Page {updated_data.get('page')}, Nom: {updated_data.get('custom_name', 'Sans nom')}")
        
        try:
            self.highlights_grid.update_highlight(updated_data)
            self._save_to_extraction_file()
            print(f"SUCCESS: Highlight sauvegarde: {updated_data.get('custom_name', 'Sans nom')}")
            
        except Exception as e:
            print(f"ERREUR: Erreur sauvegarde highlight: {e}")
            self._force_refresh_all_cards()
    
    def _on_edit_cancelled(self):
        """Callback quand l'edition est annulee."""
        print("INFO: Edition annulee")
    
    def _on_highlight_deleted(self, highlight_data):
        """Callback quand un highlight est supprime depuis le panneau d'edition."""
        print(f"DEBUG: Suppression highlight: Page {highlight_data.get('page')}, Nom: {highlight_data.get('custom_name', 'Sans nom')}")
        
        # CORRECTION: Nettoyer la selection si c'est la fiche selectionnee qui est supprimee
        if (self.selected_card and 
            self.selected_card.highlight_data.get('page') == highlight_data.get('page') and
            self.selected_card.highlight_data.get('text', '')[:50] == highlight_data.get('text', '')[:50]):
            self.selected_card = None
            self.edit_panel._show_default_message()
        
        try:
            # Supprimer de la grille
            self.highlights_grid._on_highlight_deleted(highlight_data)
            self._update_highlights_count()
            self._save_to_extraction_file()
            print(f"SUCCESS: Highlight supprime: {highlight_data.get('custom_name', 'Sans nom')}")
            
        except Exception as e:
            print(f"ERREUR: Erreur suppression highlight: {e}")
    
    def _force_refresh_all_cards(self):
        """Force le rafraichissement de toutes les cartes."""
        try:
            for i, (card, data) in enumerate(zip(self.highlights_grid.cards, self.highlights_grid.highlights_data)):
                card.update_data(data)
            print("INFO: Toutes les cartes ont ete rafraichies")
        except Exception as e:
            print(f"ERREUR: Erreur rafraichissement: {e}")
    
    # Gestion des fichiers
    
    def _save_to_extraction_file(self):
        """Sauvegarde les modifications dans le fichier d'extraction."""
        if not self.extraction_file_path or not os.path.exists(self.extraction_file_path):
            self._find_or_create_extraction_file()
        
        if self.extraction_file_path:
            try:
                highlights_data = self.highlights_grid.get_highlights_data()
                
                content = "=== HIGHLIGHTS KINDLE EXTRAITS ===\n"
                content += f"Genere le: {datetime.now().strftime('%d/%m/%Y a %H:%M')}\n"
                content += f"Nombre total: {len(highlights_data)}\n\n"
                
                for i, highlight in enumerate(highlights_data, 1):
                    title = highlight.get('custom_name', f"Page {highlight.get('page', '?')}")
                    content += f"--- {i}. {title} ---\n"
                    content += f"Page: {highlight.get('page', '?')}\n"
                    content += f"Confiance: {highlight.get('confidence', 0):.1f}%\n"
                    
                    if highlight.get('modified'):
                        content += f"Modifie le: {highlight.get('modified_date', 'N/A')}\n"
                    
                    content += f"\nTexte:\n{highlight.get('text', '')}\n\n"
                    content += "-" * 50 + "\n\n"
                
                with open(self.extraction_file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                print(f"INFO: Fichier d'extraction mis a jour: {self.extraction_file_path}")
                
            except Exception as e:
                print(f"ERREUR: Erreur lors de la sauvegarde: {e}")
    
    def _find_or_create_extraction_file(self):
        """Trouve ou cree le fichier d'extraction."""
        extractions_dir = "extractions"
        if os.path.exists(extractions_dir):
            txt_files = [f for f in os.listdir(extractions_dir) if f.endswith('.txt')]
            if txt_files:
                latest_file = max(txt_files, key=lambda f: os.path.getctime(os.path.join(extractions_dir, f)))
                self.extraction_file_path = os.path.join(extractions_dir, latest_file)
                return
        
        if not os.path.exists(extractions_dir):
            os.makedirs(extractions_dir)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"highlights_extraits_{timestamp}.txt"
        self.extraction_file_path = os.path.join(extractions_dir, filename)
    
    # Export Word - VERSION CORRIGEE SANS EMOJIS
    
    def _on_export_word_clicked(self):
        """Exporte tous les highlights vers un document Word - VERSION PROFESSIONNELLE."""
        print("DEBUG: Bouton Export Word clique - Debut de la fonction")
        
        if not WORD_AVAILABLE:
            print("ERREUR: python-docx non disponible")
            messagebox.showerror(
                "Export Word indisponible", 
                "La bibliotheque python-docx n'est pas installee.\n\n" +
                "Installez-la avec: poetry add python-docx\n" +
                "Puis redemarrez l'application."
            )
            return
        
        highlights_data = self.highlights_grid.get_highlights_data()
        print(f"DEBUG: Nombre de highlights trouves: {len(highlights_data)}")
        
        if not highlights_data:
            messagebox.showwarning("Aucun highlight", "Aucun highlight a exporter.")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="Exporter vers Word pour Zotero",
            defaultextension=".docx",
            filetypes=[("Documents Word", "*.docx"), ("Tous les fichiers", "*.*")],
            initialfile=f"highlights_kindle_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        
        print(f"DEBUG: Fichier selectionne: {file_path}")
        
        if file_path:
            try:
                doc = Document()
                
                title = doc.add_heading('Highlights Kindle - Extraits AllamBik', 0)
                title.alignment = 1
                
                info_para = doc.add_paragraph()
                info_para.add_run("Document genere le: ").bold = True
                info_para.add_run(f"{datetime.now().strftime('%d/%m/%Y a %H:%M')}\n")
                info_para.add_run("Nombre total d'extraits: ").bold = True
                info_para.add_run(f"{len(highlights_data)} highlights\n")
                info_para.add_run("Source: ").bold = True
                info_para.add_run("Application AllamBik v3.0\n")
                info_para.add_run("Compatible: ").bold = True
                info_para.add_run("Zotero, Obsidian, Notion")
                
                doc.add_paragraph("=" * 60)
                
                print(f"DEBUG: Traitement de {len(highlights_data)} highlights...")
                
                for i, highlight in enumerate(highlights_data, 1):
                    print(f"   Processing highlight {i}/{len(highlights_data)}")
                    
                    title_text = highlight.get('custom_name', f"Extrait Page {highlight.get('page', '?')}")
                    heading = doc.add_heading(f"{i}. {title_text}", level=2)
                    
                    meta_para = doc.add_paragraph()
                    meta_para.add_run("Page: ").bold = True
                    meta_para.add_run(f"{highlight.get('page', '?')}")
                    
                    meta_para.add_run("  |  Confiance: ").bold = True
                    meta_para.add_run(f"{highlight.get('confidence', 0):.1f}%")
                    
                    if highlight.get('modified'):
                        meta_para.add_run("  |  Modifie le: ").bold = True
                        try:
                            mod_date = datetime.fromisoformat(highlight.get('modified_date', ''))
                            meta_para.add_run(mod_date.strftime('%d/%m/%Y a %H:%M'))
                        except:
                            meta_para.add_run(highlight.get('modified_date', 'N/A'))
                    
                    text_para = doc.add_paragraph()
                    text_para.add_run("Contenu: ").bold = True
                    
                    content = highlight.get('text', '').strip()
                    if content:
                        text_content = doc.add_paragraph(content)
                        text_content.style = 'Quote'
                    else:
                        doc.add_paragraph("[Aucun texte]")
                    
                    if highlight.get('custom_name'):
                        tag_para = doc.add_paragraph()
                        tag_para.add_run("Tags: ").bold = True
                        tag_para.add_run(f"kindle, highlights, {highlight.get('custom_name', '').lower()}")
                    
                    if i < len(highlights_data):
                        doc.add_paragraph("-" * 40)
                
                doc.add_page_break()
                footer = doc.add_paragraph()
                footer.add_run("Document genere par AllamBik v3.0\n").bold = True
                footer.add_run("Compatible avec Zotero, Obsidian, Notion\n")
                footer.add_run(f"Export effectue le {datetime.now().strftime('%d/%m/%Y a %H:%M')}")
                
                print("DEBUG: Sauvegarde du document...")
                doc.save(file_path)
                
                print("SUCCESS: Export Word reussi!")
                messagebox.showinfo(
                    "Export Word reussi", 
                    f"Document Word cree avec succes!\n\n" +
                    f"Fichier: {os.path.basename(file_path)}\n" +
                    f"Emplacement: {os.path.dirname(file_path)}\n\n" +
                    f"{len(highlights_data)} highlights exportes\n" +
                    f"Pret pour import dans Zotero!"
                )
                
            except Exception as e:
                print(f"ERREUR: Erreur lors de l'export Word: {e}")
                messagebox.showerror(
                    "Erreur d'export Word", 
                    f"Erreur lors de la creation du document Word:\n\n{str(e)}\n\n" +
                    "Verifiez que:\n" +
                    "- python-docx est installe (poetry add python-docx)\n" +
                    "- Le fichier de destination est accessible\n" +
                    "- Vous avez les droits d'ecriture"
                )
    
    # Autres handlers
    
    def _on_search_changed(self, event):
        """Gere les changements de recherche avec filtrage et surbrillance."""
        search_text = self.search_entry.get().lower().strip()
        self.current_search = search_text
        
        if not search_text:
            # Afficher toutes les fiches si recherche vide
            self._show_all_cards()
        else:
            # Filtrer et surligner les fiches correspondantes
            self._filter_and_highlight_cards(search_text)
        
        self._update_highlights_count()
    
    def _show_all_cards(self):
        """Affiche toutes les fiches sans filtrage."""
        for card in self.highlights_grid.cards:
            try:
                card.pack_info()  # Verifier si la carte est affichee
                card.pack(fill="x", padx=5, pady=5)
            except:
                # Si pas en mode pack, utiliser grid
                try:
                    row = self.highlights_grid.cards.index(card) // self.highlights_grid.columns
                    col = self.highlights_grid.cards.index(card) % self.highlights_grid.columns
                    card.grid(row=row, column=col, padx=5, pady=5, sticky="ew")
                except:
                    pass
            
            # Retirer le surlignage de recherche
            card._remove_search_highlight()
    
    def _filter_and_highlight_cards(self, search_text):
        """Filtre et surligne les fiches selon la recherche."""
        visible_count = 0
        
        for i, card in enumerate(self.highlights_grid.cards):
            highlight_data = card.highlight_data
            
            # Verifier si la recherche correspond
            text_match = search_text in highlight_data.get('text', '').lower()
            name_match = search_text in highlight_data.get('custom_name', '').lower()
            page_match = search_text in str(highlight_data.get('page', ''))
            
            if text_match or name_match or page_match:
                # Afficher la carte
                try:
                    row = visible_count // self.highlights_grid.columns
                    col = visible_count % self.highlights_grid.columns
                    card.grid(row=row, column=col, padx=5, pady=5, sticky="ew")
                    visible_count += 1
                    
                    # Ajouter surlignage
                    card._add_search_highlight(search_text)
                    
                except Exception as e:
                    print(f"ERREUR: Affichage carte {i}: {e}")
            else:
                # Cacher la carte
                try:
                    card.grid_remove()
                    card._remove_search_highlight()
                except:
                    pass
    
    def _clear_search(self):
        """Efface la recherche et affiche toutes les fiches."""
        self.search_entry.delete(0, "end")
        self.current_search = ""
        self._show_all_cards()
        self._update_highlights_count()
    
    def _update_highlights_count(self):
        """Met a jour le compteur de highlights."""
        count = self.highlights_grid.get_count()
        self.highlights_title.configure(text=f"HIGHLIGHTS EXTRAITS ({count})")
    
    # Handlers existants
    
    def _on_zone_selected(self, zone: Tuple[int, int, int, int]):
        """Callback quand une zone est selectionnee."""
        if hasattr(self.viewmodel, 'custom_scan_zone'):
            self.viewmodel.custom_scan_zone = zone
        else:
            self.viewmodel.custom_scan_zone = zone
        
        x, y, w, h = zone
        print(f"INFO: Zone de scan definie: {w}x{h} pixels a la position ({x},{y})")
    
    def _on_start_clicked(self):
        """Gere le clic sur Demarrer."""
        # Utiliser les pages detectees si disponibles
        if hasattr(self.viewmodel, 'detected_pages') and self.viewmodel.detected_pages:
            total_pages = self.viewmodel.detected_pages
            print(f"INFO: Utilisation automatique de {total_pages} pages detectees")
            
            if self.async_loop:
                asyncio.run_coroutine_threadsafe(
                    self.viewmodel.start_extraction_command(total_pages),
                    self.async_loop
                )
            return
        
        # Sinon demander manuellement
        dialog = ctk.CTkInputDialog(
            text="Nombre total de pages du livre:",
            title="Configuration"
        )
        pages_str = dialog.get_input()
        
        if pages_str:
            try:
                total_pages = int(pages_str)
                if total_pages > 0:
                    if self.async_loop:
                        asyncio.run_coroutine_threadsafe(
                            self.viewmodel.start_extraction_command(total_pages),
                            self.async_loop
                        )
            except ValueError:
                pass
    
    def _on_stop_clicked(self):
        """Gere le clic sur Arreter."""
        if self.async_loop:
            asyncio.run_coroutine_threadsafe(
                self.viewmodel.stop_extraction_command(),
                self.async_loop
            )
    
    # Callbacks du ViewModel
    
    def _on_detect_pages_clicked(self):
        """Lance la detection automatique du nombre de pages."""
        print("INFO: Detection automatique demandee")
        
        from tkinter import messagebox
        
        # Verifier si le detecteur existe
        if not hasattr(self.viewmodel, 'page_detector') or self.viewmodel.page_detector is None:
            print("INFO: Detecteur non configure")
            messagebox.showinfo(
                "Detection non disponible",
                "La detection automatique n'est pas encore configuree.\n\n" +
                "Vous pouvez entrer le nombre de pages manuellement."
            )
            return
        
        # Demander confirmation
        if messagebox.askyesno("Detection des pages", "Lancer la detection automatique du nombre de pages?"):
            print("INFO: Detection acceptee par l'utilisateur")
            self.detect_button.configure(state="disabled", text="DETECTION EN COURS...")
            
            # Si detect_pages_command existe
            if hasattr(self.viewmodel, 'detect_pages_command') and self.async_loop:
                import asyncio
                
                def on_complete(future):
                    self.detect_button.configure(state="normal")
                    try:
                        result = future.result()
                        if self.viewmodel.detected_pages:
                            self.detect_button.configure(
                                text=f"{self.viewmodel.detected_pages} PAGES",
                                fg_color="#00aa00"
                            )
                            print(f"SUCCESS: {self.viewmodel.detected_pages} pages detectees")
                    except Exception as e:
                        print(f"ERREUR: {e}")
                        self.detect_button.configure(text="DETECTER NOMBRE DE PAGES")
                
                future = asyncio.run_coroutine_threadsafe(
                    self.viewmodel.detect_pages_command(),
                    self.async_loop
                )
                future.add_done_callback(lambda f: self.after(0, lambda: on_complete(f)))
            else:
                # Fallback: simulation simple
                print("INFO: Mode test - simulation de detection")
                messagebox.showinfo("Test", "Le detecteur sera bientot fonctionnel!")
                self.detect_button.configure(state="normal", text="DETECTER NOMBRE DE PAGES")
    
    def _on_import_json_clicked(self):
        """Charge un fichier JSON d'extraction precedente."""
        from tkinter import filedialog, messagebox
        import json
        import os
        
        print("INFO: Import JSON demande")
        
        # Ouvrir le dialogue de selection de fichier
        try:
            file_path = filedialog.askopenfilename(
                title="Selectionner un fichier JSON d'extraction",
                initialdir=os.path.abspath("extractions") if os.path.exists("extractions") else os.getcwd(),
                filetypes=[
                    ("Fichiers JSON", "*.json"),
                    ("Tous les fichiers", "*.*")
                ]
            )
            
            if not file_path:
                print("Selection annulee par l'utilisateur")
                return
                
        except Exception as e:
            print(f"Erreur dialogue de fichier: {e}")
            messagebox.showerror("Erreur", f"Impossible d'ouvrir le dialogue de fichier:\n{str(e)}")
            return
        
        # Charger le fichier selectionne
        try:
            print(f"Chargement de: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Extraire les highlights du JSON
            highlights = []
            if isinstance(data, list):
                highlights = data
            elif isinstance(data, dict):
                highlights = data.get('highlights', data.get('results', []))
            
            # Ajouter chaque highlight a l'interface
            count = 0
            for i, h in enumerate(highlights):
                if isinstance(h, dict):
                    # Creer un dictionnaire avec tous les champs necessaires
                    highlight_data = {
                        'page': h.get('page', h.get('page_number', i // 3 + 1)),
                        'text': h.get('text', h.get('extracted_text', '')),
                        'confidence': h.get('confidence', h.get('confidence_score', 85)),
                        'timestamp': h.get('timestamp', "2024-01-01T00:00:00"),
                        'source_image': h.get('source_image', None),
                        'coordinates': h.get('coordinates', None),
                        'validated': h.get('validated', False),
                        'modified': h.get('modified', False)
                    }
                    
                    # Ajouter directement via highlights_grid
                    self.highlights_grid.add_highlight(highlight_data)
                    count += 1
            
            # Mettre a jour le compteur d'highlights
            self._update_highlights_count()
            
            print(f"[OK] {count} highlights charges du JSON")
            
            messagebox.showinfo(
                "Import reussi",
                f"{count} highlights charges depuis:\n{os.path.basename(file_path)}"
            )
            
        except Exception as e:
            print(f"Erreur import JSON: {e}")
            import traceback
            traceback.print_exc()
            messagebox.showerror("Erreur", f"Impossible de charger le fichier:\n{str(e)}")
    
    def _on_state_changed(self, state: ViewState):
        """Met a jour l'UI selon l'etat."""
        def update():
            self.start_button.configure(
                state="normal" if self.viewmodel.can_start else "disabled"
            )
            self.stop_button.configure(
                state="normal" if self.viewmodel.can_stop else "disabled"
            )
            
            has_highlights = self.highlights_grid.get_count() > 0
            self.export_word_button.configure(
                state="normal" if has_highlights else "disabled"
            )
            
            if state == ViewState.ERROR:
                self.progress_circle.configure(fg_color=self.colors['error'])
        
        self._schedule_update(update)
    
    def _on_progress_changed(self):
        """Met a jour la progression."""
        def update():
            self.progress_circle.update_progress(
                current=self.viewmodel.current_progress,
                phase1=self.viewmodel.phase1_progress,
                phase2=self.viewmodel.phase2_progress
            )
            self.progress_label.configure(text=self.viewmodel.progress_message)
            
            self.stats_labels["pages_scanned"].configure(
                text=str(self.viewmodel.pages_scanned)
            )
            self.stats_labels["pages_with_content"].configure(
                text=str(self.viewmodel.pages_with_content)
            )
            self.stats_labels["highlights_count"].configure(
                text=str(self.viewmodel.highlights_count)
            )
        
        self._schedule_update(update)
    
    def _on_highlight_added(self, highlight: HighlightViewModel):
        """Ajoute un highlight a la grille."""
        def update():
            highlight_data = {
                'page': highlight.page,
                'text': highlight.text,
                'confidence': highlight.confidence,
                'timestamp': datetime.now().isoformat(),
                'source_image': None,
                'coordinates': None,
                'validated': False,
                'modified': False
            }
            
            self.highlights_grid.add_highlight(highlight_data)
            self._update_highlights_count()
            self._save_to_extraction_file()
        
        self._schedule_update(update)